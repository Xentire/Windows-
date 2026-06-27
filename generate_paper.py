"""
Generate the course paper docx file - personal style edition.
"""
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Pt(21)

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1: run.font.size = Pt(14)
        elif level == 2: run.font.size = Pt(12)
        elif level == 3: run.font.size = Pt(10.5)
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    if bold: run.bold = True
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for line in code_text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

# ============================
# TITLE
# ============================
for _ in range(6):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('Windows程序设计课程作业')
title_run.font.name = '黑体'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_run.font.size = Pt(22)
title_run.bold = True

doc.add_paragraph()
subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_p.add_run('基于WinForms的图像处理工具设计与实现')
subtitle_run.font.name = '黑体'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
subtitle_run.font.size = Pt(16)

for _ in range(6):
    doc.add_paragraph()

for item in ['学    院：数学与统计学院', '专    业：', '姓    名：', '学    号：', '日    期：2026年6月']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(item)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

doc.add_page_break()

# ============================
# 1. 课程综述
# ============================
add_heading_text('1. 课程综述', level=1)
add_heading_text('1.1 课程重点技术总结', level=2)

# 1.1.1 C#
add_heading_text('1.1.1 C#基础语法', level=3)
add_para('这学期从零开始接触C#。之前只在大一学过一点C语言，属于那种能看懂的层次，真让自己写就犯怵。C#给我的第一印象是——它把很多C/C++里容易踩坑的地方替你挡掉了。比如内存管理，以前写C的时候malloc/free搞不清楚就内存泄漏，C#有GC，虽然不是说可以完全不操心资源释放，但至少大部分时候你不用盯着每一处new去配delete。')
add_para('学面向对象那部分的时候，我其实一开始挺懵的。封装、继承、多态，课本上的例子看懂了，一动手就不知道怎么拆分类。后来是在实验课上做那个学生管理系统的练习，把Student类和Course类分开，再写一个GradeManager来关联它们，才慢慢体会到"一个类只干一件事"的道理。这个习惯在后面做图像处理项目时帮了大忙——ImageProcessor就只负责算法，文件读写交给ImageFileHandler，UI的事归MainForm管。各管各的，改一个地方不用操心另一个地方崩溃。')
add_para('委托和事件这块花了我不少时间理解。说实话，刚学的时候觉得这语法好奇怪——为什么要把方法当参数传来传去？直到写WinForms界面，给按钮绑点击事件：btn.Click += (s, e) => { ... }，才突然明白这就是委托啊。lambda表达式写起来确实比单独声明一个方法再绑定简洁太多了。项目中所有菜单项和工具栏按钮都用lambda绑的，我个人感觉这是C#里用得最舒服的特性之一。')
add_para('async/await也是这学期新学的东西。以前写程序都是同步的，一行执行完才走下一行。第一次遇到界面卡死的问题是在做网络编程实验的时候——主线程上发了个HTTP请求，等响应的过程中整个窗口拖都拖不动。后来理解了消息循环（message loop）这个东西：WinForms的UI线程就是在不断处理消息队列，你把它堵住了，界面当然动不了。Task.Run把重活扔给后台线程，配合IProgress报告进度，这个模式在这次图像处理项目里反复用到。')

# 1.1.2 WinForm
add_heading_text('1.1.2 WinForm窗体与控件', level=3)
add_para('WinForms给我的感觉是——它不像现在那些前端框架那么花哨，控件长得就是Windows原生风格，但是上手快，搭一个能用的工具软件效率很高。课程从最简单的Form和Button开始，逐步讲了容器控件（Panel、GroupBox、TabControl）、菜单工具栏（MenuStrip、ToolStrip、ContextMenuStrip）、数据显示控件（DataGridView、ListView、TreeView）和对话框（OpenFileDialog、SaveFileDialog、ColorDialog等）。')
add_para('做这个图像处理项目的时候，MainForm的整个界面是我自己一行行代码搭的，没依赖设计器。一开始觉得手写布局好麻烦，各种Dock、Anchor、Padding调来调去。但写完之后反而觉得手写比拖控件更可控——你知道每个控件的定位逻辑是什么，出了问题也好排查。比如那个PictureBox要配合Panel实现滚动查看大图，核心就是两点：PictureBox设AutoSize让它的尺寸跟图片一样大，Panel设AutoScroll让它在内容超出可见区域时自动出滚动条。就这么两行属性的设置，搞懂了原理其实很简单。')
add_para('StatusStrip的状态信息展示我个人觉得对用户体验提升特别大。没状态栏的时候，你打开一张图片，除了看到图本身，什么都不知道——多大的文件？什么格式？多少像素？加了状态栏之后，一打开图，底下直接显示"1920x1080 | JPEG | 2.3 MB"，心里就有数了。处理进度条也很实用，加个大半径模糊的时候可能要等几秒钟，有个进度条就不至于让用户以为程序卡死了。')

# 1.1.3 并发
add_heading_text('1.1.3 并发控制', level=3)
add_para('并发这块说实话学得比较吃力。线程安全、锁、死锁这些概念光看书根本搞不明白，非得自己踩过坑才有感觉。印象最深的一次是做一个多线程下载的实验，开了5个线程同时写一个共享的List<int>，跑着跑着突然抛了个IndexOutOfRangeException，排查半天发现是两个线程同时执行Add操作导致内部数组扩容时数据错乱了。换成ConcurrentBag之后就再没出过这个问题。')
add_para('图像处理项目里的并发倒没那么复杂——主要是单向的：后台线程处理图像，处理完了把结果Bitap传回UI线程显示。Progress<T>这个类设计得很妙，你不用手动Invoke或者BeginInvoke，它内部帮你把SynchronizationContext捕获好了，Report回调自动在正确的线程上执行。省了不少事。')

# 1.1.4 网络 + 文件
add_heading_text('1.1.4 网络连接与文件处理', level=3)
add_para('网络编程部分做了TCP聊天程序的实验，虽然这次图像处理工具没用到网络功能，但Socket编程的思路——建立连接、收发数据、关闭连接——对理解C/S架构帮助很大。HttpClient封装得比原始Socket好用太多了，一个GetAsync就能拿到响应，感觉微软在这块的API设计上下了功夫。')
add_para('文件处理在项目里是实实在在用到的。ImageFileHandler的OpenImage方法没有直接调Bitmap.FromFile()，而是先把文件读到MemoryStream再从流创建Bitmap。为什么这么绕一下？因为Bitmap.FromFile()会一直锁着源文件，你想保存覆盖的时候就报"文件被占用"。用MemoryStream中转，读完马上释放文件句柄，后面再保存就不会有问题了。这个小细节是在StackOverflow上看到的，自己遇到之前完全想不到。')

# 1.1.5 图像处理
add_heading_text('1.1.5 图像处理技术', level=3)
add_para('图像处理是这学期我觉得最有意思的部分，也是最后选来做期末项目的方向。课堂上讲了图像的基础表示——每个像素就是几个byte，灰度图一个byte，彩色图三个byte（RGB）或四个（ARGB）。听起来简单，但正是这个简单的模型支撑了所有复杂的图像算法。')
add_para('GetPixel/SetPixel是入门用的，做做实验还行。真处理一张几百万像素的照片，这两个方法慢到你怀疑人生。为啥慢？因为每次调用GetPixel都要跨越托管代码和非托管GDI+代码之间的边界，这个"跨域"开销累积起来不得了。老师课上演示了LockBits的做法——直接把整个位图的内存区域锁住，拿到一个指针，然后你就可以像操作数组一样遍历像素了。配合unsafe代码块的指针运算，处理一张1920x1080的图从十来分钟降到一秒以内，这个对比太直观了。这也是我第一次写unsafe代码，一开始挺紧张，怕把内存搞崩了，后来发现只要搞清楚stride（每行的字节跨度，可能因为4字节对齐而大于width*3）和BGR的字节顺序，就没那么容易出错。')
add_para('卷积那节课给我印象很深。之前总听说"模糊滤镜""锐化滤镜"这些词，但不知道背后就是简单的加权求和。一个3x3的小矩阵在图像上滑过去，每个位置的像素值用它和周围邻居的加权平均替代。模糊核就是大家都一样权重然后归一化，锐化核是中间权重大、旁边权重为负。Sobel边缘检测更有意思——用两个核分别算水平和垂直方向的梯度，然后勾股定理合成幅度。数学课学的矩阵和向量在这里全都用上了。')

# 1.1.6 数据库
add_heading_text('1.1.6 数据库连接与处理', level=3)
add_para('数据库部分做了SQL Server的实验，从建表、写SQL查询到ADO.NET读写数据。参数化查询防止SQL注入这一点印象很深——之前不知道直接把用户输入拼到SQL语句里有多大风险，后来看了演示用的SQL注入字符串才意识到问题的严重性。图像处理工具项目本身没用到数据库（定位是个轻量工具，文件系统够用了），但如果要扩展的话，比如记录用户的处理历史或者保存常用的滤镜参数配置，SQLite应该是个挺好的选择——不需要安装数据库服务，一个文件搞定。')

add_para('回顾这学期学的内容，我觉得最大的收获倒不是记住了多少API或者控件的名字，而是一种"遇到问题知道往哪个方向去找答案"的能力。比如要做一个图像处理功能，我知道核心在System.Drawing命名空间里，要处理Bitmap的像素数据LockBits是最快的，要异步就用Task.Run配Progress<T>。这种知识框架比具体的语法细节重要得多。', bold=False)

# 1.2 深入实践
add_heading_text('1.2 技术深入实践——Sobel边缘检测的实现', level=2)
add_para('选择图像处理作为深入方向，纯粹是因为兴趣。滤镜效果是直观可见的——你写对了，图像真的会变；你写错了，出来的结果是乱码或者全黑一片，反馈特别直接。不像有些方向（比如网络通信），出了bug可能半天找不到在哪。而且图像处理涉及的知识面很广，写一个边缘检测算法，C#基础语法、unsafe指针、数学公式、单元测试全都用上了，很适合拿来检验一学期的学习成果。')

add_heading_text('1.2.1 原理简述', level=3)
add_para('Sobel算子的原理其实不复杂。想象图像是一个亮度函数f(x,y)，边缘就是函数值剧烈变化的地方。怎么找到这些地方？求导。对离散的图像来说，导数就是相邻像素的差值。Sobel用两个3x3的卷积核来近似x方向和y方向的导数：Gx核检测垂直边缘（因为它是从左边减右边），Gy核检测水平边缘。两个方向的梯度合成后，取欧几里得距离sqrt(Gx²+Gy²)就是最终的边缘强度。')
add_code_block('''Gx = [[-1, 0, 1],     Gy = [[-1, -2, -1],
      [-2, 0, 2],           [ 0,  0,  0],
      [-1, 0, 1]]           [ 1,  2,  1]]''')
add_para('注意Gy核中间三列全是0——对于水平方向的变化，它完全不做响应，只对垂直方向有输出。反过来Gx的中间行全是0。两个核各司其职，最后的sqrt操作把它们统一起来。')

add_heading_text('1.2.2 实现过程', level=3)
add_para('实现的时候主要花心思在两个地方：性能和对边界像素的处理。性能方面，全程走LockBits + unsafe路线，把源位图的数据一次性拷到一个byte数组里，在数组上做卷积计算，最后再拷回目标位图。边界像素的邻域会越界，我选了clamp策略——把越界的坐标拉到最近的边界上，等价于把边缘像素的颜色向外复制。这个策略跟OpenCV的BORDER_REPLICATE是一样的，不会在图像边缘产生假的响应。')
add_para('还有一个细节：我选择对RGB三个通道分别做梯度，然后在每个像素上合成，而不是先转灰度再检测。这样做的好处是能保留颜色变化产生的边缘。假如一张图里红色块贴着绿色块，亮度一样，转灰度之后它们就分不出来了；但分通道做的话，红色通道里有变化、绿色通道里也有变化，sqrt合成后边缘照样能被检测到。多花了点计算量，但效果确实更好。')

add_code_block('''// 核心循环 - 对每个像素的3x3邻域卷积
for (int y = 0; y < h; y++)
{
    for (int x = 0; x < w; x++)
    {
        double gxR = 0, gxG = 0, gxB = 0;
        double gyR = 0, gyG = 0, gyB = 0;
        for (int ky = -1; ky <= 1; ky++)
        {
            for (int kx = -1; kx <= 1; kx++)
            {
                int sx = Math.Clamp(x + kx, 0, w - 1);
                int sy = Math.Clamp(y + ky, 0, h - 1);
                int idx = sy * stride + sx * 3;
                double gx = gxKernel[ky+1, kx+1];
                double gy = gyKernel[ky+1, kx+1];
                gxB += srcBuf[idx]   * gx; gyB += srcBuf[idx]   * gy;
                gxG += srcBuf[idx+1] * gx; gyG += srcBuf[idx+1] * gy;
                gxR += srcBuf[idx+2] * gx; gyR += srcBuf[idx+2] * gy;
            }
        }
        int di = y * stride + x * 3;
        dstBuf[di]   = ClampToByte((int)Math.Sqrt(gxB*gxB + gyB*gyB));
        dstBuf[di+1] = ClampToByte((int)Math.Sqrt(gxG*gxG + gyG*gyG));
        dstBuf[di+2] = ClampToByte((int)Math.Sqrt(gxR*gxR + gyR*gyR));
    }
    progress?.Report((y + 1) * 100 / h);
}''')

add_heading_text('1.2.3 测试验证', level=3)
add_para('测试这块我写的比较认真，因为图像算法太容易出那种"看起来差不多"但实际上是错的bug。比如灰度公式里面三个通道的权重，如果你把R和B的系数搞反了，输出还是一张灰度图，肉眼看不出毛病，但数值全偏了。所以我写单元测试的时候，不是随便跑跑看看结果就完事，而是针对每种情况算了期望值再跟实际比对。')
add_para('均匀图像测边缘检测：纯色图没有任何亮度变化，经过Sobel之后所有像素应该接近0。我建了一张20x20全灰色图，处理完取中间像素看，结果R<10，跟预期一致。另一半黑一半白的图，在分界线x=5的位置，梯度响应跳到R>100——这个跳跃正好说明算法确实检测到了边缘。刚开始做的时候有一个错误，就是Sobel算出来的梯度值可能超过255，我在写回byte之前忘了做clamp，导致溢出后出现了奇怪的黑色条纹。加上ClampToByte之后就正常了。')

# ============================
# 2. UI设计
# ============================
add_heading_text('2. UI设计样例说明', level=1)

add_para('UI设计是这门课要求自学的部分。之前我对"界面好看"的理解基本就是找几个图标放上去，颜色选亮一点就行。后来看了一些博客和视频教程才意识到，UI设计远不止审美——它首先是信息架构的问题：用户一眼看过去能不能找到他想要的？操作路径是不是最短的？当前状态是不是一目了然的？')

add_heading_text('2.1 参考的学习资源', level=2)
add_para('推荐几个我觉得确实有帮助的资源：')
add_para('一是微软官方的Fluent Design文档（https://fluent2.microsoft.design/）。这不是那种泛泛而谈的设计原则，而是有很具体的规范——按钮间距多少像素、何时用图标配文字何时纯图标就行、不同层级的标题字号怎么定。我项目里工具栏用纯文字按钮（DisplayStyle = Text）而不是图标，就是参考了它说的"当功能超过5个时，文字比图标更不易让用户困惑"。')
add_para('二是Refactoring UI这个网站（https://www.refactoringui.com/），两个Tailwind CSS的作者做的。虽然它主要是给Web设计的建议，但很多原则桌面应用也适用：比如"层级不是靠颜色而是靠大小来区分的"——重要信息大一点、次要信息灰一点小一点，这种对比天然引导眼睛的移动轨迹。')
add_para('三是B站上"设计早知道"UP主的WinForms美化教程。有一集讲怎么用GDI+自己画ToolStrip的背景（重写ToolStripRenderer），用LinearGradientBrush拉一个从浅灰到深灰的渐变，比默认的纯灰好看很多，还不用引入第三方库。不过我这个项目考虑到开发时间有限，没有完全自定义绘制，还是用的原生ToolStrip风格——功能优先吧。')

add_heading_text('2.2 本项目UI设计思路', level=2)
add_para('我的界面设计遵循了几个简单的原则。第一是功能分层：最常用的操作——打开、保存、撤销、重做、常用滤镜——放在工具栏上直接可见；不那么常用的（比如调整亮度的具体数值）收在菜单里，点进去弹出对话框。这不是我发明的，Windows的资源管理器、Office软件都这么干的，用户早就习惯了这种模式。')
add_para('第二是状态可见。打开一张图之后，状态栏立刻显示"宽x高 | 格式 | 文件大小"，不用你去右击属性查看。滤镜处理中的时候，状态栏右端出现一个进度条和一个百分比文字——用户知道系统正在工作，不是死了。')
add_para('第三是容错。所有操作都支持Ctrl+Z撤销，包括滤镜、调整、旋转。这个在专业软件（比如Photoshop）里是标配，但自己做项目的时候很多人会忽略。实现起来不复杂——就是每次操作前把当前Bitmap存一份副本——但用户体验的天差地别。')
add_para('关于配色，图像显示区的背景我设成了深灰色（RGB 40,40,40），跟大多数看图软件的暗色背景一致。不是因为喜欢暗色，而是因为：在白色背景上看亮色图片，瞳孔要适应背景的亮，图像的细节对比度就降低了；暗色背景把环境光压下去，图片本身的颜色和层次感会更突出。这是一个视觉感知上的小技巧。')

# ============================
# 3. AI辅助编程实践
# ============================
add_heading_text('3. 应用开发——AI辅助编程实践', level=1)

add_heading_text('3.1 功能模块与核心代码', level=2)

add_para('这个图像处理工具虽然不大，但五脏俱全——文件IO、UI交互、算法计算、状态管理、异步处理都有了。整个项目可以分成四个模块，各司其职：')

add_para('ImageProcessor：所有图像处理算法的集合，设计成纯静态方法。这么做有个好处——每个方法输入Bitmap输出Bitmap，不依赖外部状态，所以可以单独拿出来测试，也可以在任何地方调用。核心的像素遍历逻辑抽取到了ApplyPixelFunc这个泛用方法里，接收一个Func<byte,byte,byte,(byte,byte,byte)>委托来决定每个像素怎么变换。新加一个滤镜只需要写一个lambda描述变换规则就行了，不用改遍历逻辑。卷积类（模糊、锐化、Sobel）共用ApplyConvolution框架，接收kernel矩阵和normalization参数，也是在避免重复代码。', bold=True)

add_para('UndoRedoManager：用两个Stack<Bitmap>实现。操作前把当前状态Push进undo栈，撤销时当前状态进redo栈、从undo栈Pop出上一个状态。重做的时候反向操作。执行新操作的时候要把redo栈清空——这跟Word、Photoshop的行为一致，新的修改会干掉"重做分支"。为什么选Bitmap快照而不是命令模式？因为命令模式要为每种操作写可逆命令，而模糊这种操作数学上不可逆（信息丢了）。快照方案一个缺点是用内存多——每步存一张完整位图——但20步深度、处理1920x1080的图，撑死了也就100MB左右的占用，现代桌面应用完全不是问题。', bold=True)

add_para('MainForm：把上面的模块串起来。比较有意思的是那几个泛型RunAsync方法——所有滤镜都走同一个入口，通过Func委托区分到底执行哪个算法。这样菜单按钮和工具栏按钮的点击事件就一行代码：(_,_) => ApplyFilter(ImageProcessor.Grayscale, "Grayscale")。不用给12种操作写12个几乎一样的事件处理方法。每次处理前调SetProcessingState(true)把菜单栏工具栏禁用掉，防止用户在滤镜算到一半的时候点另一个按钮——这种防御性编程省了很多潜在的bug。', bold=True)

add_para('ImageFileHandler：文件读写，没啥复杂的。唯一值得一提的是OpenImage用了MemoryStream中转而非Bitmap.FromFile()直读——避免文件被锁导致保存失败。这事不遇到不知道，遇到了查半天才明白。', bold=True)

add_heading_text('3.2 AI辅助编程实践分析', level=2)
add_heading_text('3.2.1 Prompt作品集与分析', level=3)

add_para('【案例一：项目架构设计的Prompt】', bold=True)
add_para('这是我这个项目第一次认真用Claude Code。当时需求很模糊——我就知道要做个图像处理工具，具体怎么组织代码完全没想好。我给的Prompt大意是：需要一个C# WinForms图像处理工具，要包含哪些功能，用.NET 8，要有TDD。')
add_para('Claude返回了一个很完整的技术设计，包括ImageProcessor的API列表、UndoRedoManager的双栈方案、MainForm的控件树、甚至每个滤镜的数学公式。说实话有点超出预期，本来我只是想要个大方向，结果它把详细设计都给出来了。')
add_para('但也不全是拿来就用的。比如它建议用BackgroundWorker处理异步，我觉得太老了——课程教的是async/await/Task，就把这块改掉了。它设计的BoxBlur的radius参数含义也不对——它认为radius=3生成7x7的核，而实际我想要的是radius=1就是3x3、radius=2是5x5。这些细节上还是得自己把关。整体来看，Claude帮我省掉了大概两三个小时的设计时间，但我自己调整和验证也花了一个多小时。')

add_para('【案例二：一个坑爹的测试Bug】', bold=True)
add_para('这是最有意思的一次交互。Grayscale_SolidGreen测试一直失败——灰度值返回75，我预期是150。我反复检查了公式、看了三遍LockBits代码，甚至怀疑字节顺序搞错了。折腾了快二十分钟没找出来。')
add_para('然后我问Claude："Grayscale of solid green returns 75 but I expect 150, what is wrong?" 它几乎是秒回："Color.Green in .NET is (0, 128, 0), not (0, 255, 0). Try Color.Lime." 我当场就震惊了——我在想的是内存布局、像素格式这些"底层"问题，结果就是个颜色名称的理解偏差。Color.Green确实是#008000，这是Web标准色，我一直以为是纯绿色。')
add_para('这个案例特别能体现AI的优势：它的知识覆盖面比人广。.NET里有几百个命名颜色，普通开发者不可能记住每个的RGB值。但AI"知道"。这种"冷知识类"的bug，让人排查可能半小时，AI识别只要几秒。不过话说回来，如果Claude的诊断是错的，我如果没有足够的判断力去验证，可能就会被带到另一个坑里——这大概就是用AI最需要把握的平衡。')

add_para('【案例三：手写WinForms布局代码】', bold=True)
add_para('第三个场景是生成MainForm的InitializeComponent。我在没有Visual Studio设计器的环境下开发，只能手写所有控件的创建和布局。一个完整的MainForm，从MenuStrip到StatusStrip，初始化代码大概四百多行。')
add_para('我给Claude描述了大概的控件树结构——什么菜单下有什么子项，工具栏要哪些按钮，PictureBox怎么放。它生成了一版，我跑起来发现布局不对——PictureBox没有正确Dock，状态栏被挤到右边去了。把具体的布局问题反馈给它之后，第二版修好了Dock属性。然后我自己又手动调了一些细节：背景色、内边距、字体大小。')
add_para('这个过程让我体会到，AI生成的UI代码"能用但不精致"。它知道语法、知道该设哪些属性，但对"好看"这件事没有审美——你需要反复给它feedback调细节。这种迭代式协作大概就是我理解的"人机协同编程"——人做决策和质量把控，AI干体力活。')

add_heading_text('3.2.2 AI生成代码的人机审核报告', level=3)

add_para('审核对象：Claude协助生成的BoxBlur卷积实现。整体代码框架是合理的——LockBits拿指针、双重循环遍历像素、3x3邻域加权求和——但有几个地方我做了修改：', bold=True)

add_para('第一个问题是重复计算。AI生成的代码在每次x循环里都对ky做Math.Clamp来限制采样坐标，但对于固定的ky行来说，sy在同一次y循环里是不变的。我把sy的计算提升到了ky循环的外层、x循环的外面——虽然对3x3的小核影响不大，但这是个代码习惯问题，以后万一有人把这段复制去写大核卷积呢？')

add_para('第二个是命名。AI用了全小写的kernelWidth/kernelHeight，跟C#的camelCase不统一。改成了有意义的命名，比如kCenterX/kCenterY，一看就知道是卷积核的中心坐标。变量命名这种事情说大不大，但代码是写给人看的，一致性很重要。')

add_para('第三个值得警惕的是，AI生成的代码完全没有处理stride为负的情况。绝大多数情况下stride是正的（位图数据从上往下存），但GDI+的规范并不保证这一点。有一天如果碰到一个bottom-up存储的位图，stride就是负的，直接用srcData.Stride计算偏移就挂了。我加了一个Math.Abs保护——可能一辈子用不上，但加上的成本就是一两个字符。')

add_para('回顾整个项目使用AI辅助编程的经历，我的核心感受是：AI最擅长的不是"写出最优解"，而是"快速给出一个能跑的方案"。这个方案可能有效率问题、命名问题、边界条件考虑不周全，但它至少是一个起点。人的价值在于判断——哪些地方需要优化、哪里藏着边界条件、这个设计对未来扩展是否友好。我觉得可以把AI理解成一个经验丰富但缺乏整体把控能力的初级程序员——你需要Review它的所有代码，但它的产出确实比你从零写要快。')
add_para('什么时候应该对AI的代码特别谨慎？我总结了几类：一是跟安全相关的（文件路径、权限、认证），AI不了解你的安全策略；二是性能极端敏感的地方（比如这个项目的LockBits循环），AI给出的通用写法往往不是最优；三是跟现有系统深度耦合的部分，因为AI不可能理解你的整个代码库的上下文。其他日常的、标准化的代码——比如搭个UI框架、写个CRUD方法——交给AI辅助确实能省不少时间。')

add_heading_text('3.3 运行结果截图', level=2)
add_para('（以下为运行效果描述，实际截图将插入提交版中）', indent=False)
add_para('应用主界面截图：展示完整的窗口布局——顶部菜单栏（File/Edit/Filters/Adjustments/Transform/Help）和工具栏，中央深灰背景上的图像显示区（以AutoSize模式展示原始图片，大图时自动出现滚动条），底部状态栏左侧显示"1920x1080 | JPEG | 2.3 MB"，右侧处理进度条隐藏状态。')
add_para('滤镜效果系列截图：同一张彩色照片分别应用Grayscale、Sepia、Negative、BoxBlur、Sharpen、EdgeDetectSobel六种滤镜后的对比效果。')
add_para('功能对话框截图：Brightness调整对话框（TrackBar范围-100到+100，当前值大字显示在顶部）；Resize对话框（NumericUpDown输入宽高，"保持纵横比"复选框勾选状态）。')

# ============================
# 4. 项目开发思考
# ============================
add_heading_text('4. 项目开发思考', level=1)

add_heading_text('4.1 课程知识点综合运用总结', level=2)
add_para('把项目中用到的知识点拉一张清单，其实还挺多的：')
add_para('C#基础方面——泛型集合（Stack<Bitmap>做撤销重做）、委托与Lambda（事件绑定和像素变换函数）、unsafe代码块配指针（LockBits高性能遍历）、async/await（后台处理不卡UI）、IDisposable模式（Bitmap资源释放）。说实话在写这个项目之前，很多特性我只是"会用"但不知道"什么时候该用"。做完之后清楚多了——比如什么时候走unsafe，什么时候async就够了，什么时候加锁保线程安全。')
add_para('WinForms方面——MenuStrip、ToolStrip、PictureBox、StatusStrip这些核心控件从头搭建了一遍。自己写布局代码跟拖控件最大的区别是，你必须搞清楚Dock、Anchor、AutoSize、AutoScroll这些属性的精确作用。拖控件的时候可以凭直觉调位置，手写代码就要准确理解每个属性的效果。')
add_para('图像处理方面——LockBits内存访问、BGR字节顺序、stride对齐规则、卷积核设计、颜色空间转换。特别是stride对齐这件事，卡了我好一阵子：宽度1000像素的24位图，预期每行3000字节，实际stride可能是3000也可能是3004（补齐到4的倍数）。不搞清楚这个，指针偏移全乱套。')
add_para('测试方面——36个单元测试覆盖了所有算法的正常输入、边界条件和异常场景，先写测试后写代码的TDD流程完整走了一遍。虽然一开始写测试的时候觉得"功能都没做呢写什么测试"，但写完之后发现有个巨大好处：改代码的时候跑一遍测试就知道有没有引入回归。特别是后来把Grayscale从managed改unsafe实现那次，重构完跑测试全绿，信心大增。')
add_para('不足也很明显。一是对PixelFormat的处理不够完善，现在只支持24bpp RGB输入，碰到了索引色或16bpp的图没做转换。二是UI上没有做缩放功能——用户没法放大缩小查看图像细节。三是没有做批量处理——如果有100张图要统一加滤镜，现在只能一张一张弄。这些问题在真实产品里必须解决，但作为课程项目，时间和精力有限，暂时搁置了。')

add_heading_text('4.2 项目开发总结与反思', level=2)
add_para('踩过的坑和怎么爬出来的：', bold=True)

add_para('第一个坑是BGR字节顺序。Format24bppRgb存的是B-G-R不是R-G-B，这个跟我直觉反过来的。写Grayscale的第一版，我把第一个byte当成R了，结果红色通道的值被当蓝色算，灰度值全错。排查了好久，最后是读了MSDN上PixelFormat的文档才知道这个顺序。教训就是——涉及底层数据结构的操作，不要凭直觉，一定先把官方文档里数据格式的说明看清楚。')

add_para('第二个坑前面提到了——Color.Green不是纯绿色。这严格来说不是bug，是"概念偏差"。我记得当时看到Claude指出这个问题的时候，脑子里第一个反应是"啥？还有这种事？"然后赶紧去Visual Studio里输出了一下Color.Green的RGB值，果然(0,128,0)。从此以后在代码里只要涉及命名颜色，我都会先确认一下它的准确值。')

add_para('第三个坑是Bitmap的资源释放。一个疏忽就容易内存泄漏——Bitmap底层是GDI+的非托管资源，new出来的每一个都要Dispose。特别是在UndoRedoManager里，undo栈和redo栈里存着一堆Bitmap副本，不仔细管Dispose的话，用一会儿就OOM了。我最后在UndoRedoManager的Clear方法和Dispose方法里都加了遍历Stack逐个释放的逻辑。这确实有点烦，但没办法，GDI+就是这么设计的。')

add_para('对Windows程序设计的理解：', bold=True)
add_para('做完这个项目，我对"消息驱动"有了更具体的感受。WinForms不是直接执行你的代码，而是维持着一个消息循环Application.Run()——不停地从消息队列里取Windows消息（鼠标挪了、按键了、窗口需要重绘了），分发给对应的控件处理。在这个模型下，UI线程绝对不能阻塞——你执行一个10秒的计算，这10秒内没有任何消息能被处理，窗口自然就"假死"了。所以异步编程不是一个"可选的优化"，而是GUI开发的必需品。')
add_para('另一个理解是关于分层——把算法和界面拆开。如果没有刻意去拆，很容易把图像处理的代码直接写在按钮的Click事件里——"点一个按钮，在当前图上跑个灰度算法"。功能上没毛病，但测试就没法测了（你必须启动整个窗口才能测一个算法），代码也复用不了。把ImageProcessor独立成纯静态方法，好处立竿见影：写单元测试方便，以后把这套算法移到别的项目（比如做一个命令行版本的图像处理工具）也不用动UI代码。')

add_para('AI工具在哪些环节帮助大？', bold=True)
add_para('架构设计阶段帮助最大——Claude快速给了我一个完整的技术方案，省了自己摸着石头过河的时间。代码框架生成也不错——400行的MainForm初始化，靠AI生成然后自己微调，比自己从头敲要快好几倍。但在性能优化和资源管理上，AI的帮助很有限——它给出的卷积框架是功能正确的，但指针层面的优化（比如减少冗余的Clamp调用）需要自己去抠。Bitmap的Dispose时机这种涉及"资源生命周期"的问题，AI也容易遗漏。')
add_para('总结一下就是：AI帮你把东西做出来——快。但做得好不好——还得靠你自己。')

add_heading_text('4.3 学习规划与发展方向', level=2)
add_para('看到作业里附的字节跳动Windows客户端岗位要求，我对比了一下自己的情况。', bold=True)
add_para('C#这块算是还可以——课程学了一遍，项目也做了一个，基本语法、WinForms、文件IO、异步编程都实际用过了。但岗位要求里C#只是"至少一种语言"之一，而且同样要求C/C++。我的C++水平说实话堪忧——大一的C语言课学的，基本忘差不多了，C++的面向对象、模板、智能指针、RAII这些都没系统学过。这是接下来需要优先补的短板。')
add_para('数据结构与算法方面，课程学过基础，但LeetCode还没怎么刷。看了一眼那个岗位的面经，中等难度的题占大多数——DP、二叉树、图搜索这些。应付面试的话，基本套路还是要掌握的。我计划先拿2-3个月，每天2道LeetCode，从Hot 100开始，配合《剑指Offer》的题解理解思路。不是为了刷题而刷题，确实感觉到了——写图像处理里那个卷积核的嵌套循环，时间复杂度O(w*h*k²)，如果之前对算法复杂度分析理解更透彻，优化方案可能会想得更快。')
add_para('操作系统这块是岗位里明确写的加分项。我对进程线程的区别、虚拟内存的概念、文件系统的基本结构这些有一些了解，但更底层的东西——页表是怎么工作的、上下文切换具体做了什么、锁在操作系统层面是怎么实现的——说实话还比较模糊。打算用OSTEP（Operating Systems: Three Easy Pieces）当教材啃一遍，配合xv6的实验做一下。这个可能需要3-4个月的投入。')
add_para('时间线大概是这样：暑假先集中补C++，同时每天刷题保持手感；下学期开学之后边上课边学操作系统；寒假开始投实习简历。目标是大三下学期或暑假能拿到一个Windows开发或后端开发的实习。不一定非要是字节那种大厂，先入行比什么都重要。')

# ============================
# 5. 致谢
# ============================
add_heading_text('5. 致谢', level=1)
add_para('这学期的Windows程序设计课确实花了不少时间——实验一个接一个，作业也不少。但回头看，是那种"累但值得"的课。')
add_para('感谢授课老师一个学期的讲解。说实话Windows开发在2026年不算热门方向了，大家都在追前端和AI，但老师的课让我觉得做桌面应用也挺酷的——尤其是LockBits那节课，性能从十几分钟优化到一秒，那种"原来还能这么搞"的冲击感我现在还记得。')
add_para('感谢教辅在每次实验课上的辅导和作业批改。有一次我的实验报告里代码格式一团乱，教辅在评语里一条条指出来，虽然当时觉得有点吹毛求疵，但后来确实养成了一直保持代码整洁的习惯。')
add_para('感谢班级群里一起讨论问题的同学，特别是帮我搞定stride对齐问题的那位。感谢Stack Overflow上不知名的网友，Bitmap文件锁定问题就是在那找到答案的。Claude Code在开发过程中的协助也值得感谢——定位Color.Green那个bug如果没有它，我可能还要多调试半小时。')
add_para('感谢所有在这个学期帮助过我的人。')

# ============================
# 6. 参考文献
# ============================
add_heading_text('6. 参考文献', level=1)
refs = [
    '[1] Microsoft. .NET Documentation - Windows Forms[EB/OL]. https://learn.microsoft.com/en-us/dotnet/desktop/winforms/',
    '[2] Microsoft. System.Drawing Namespace[EB/OL]. https://learn.microsoft.com/en-us/dotnet/api/system.drawing',
    '[3] Gonzalez R C, Woods R E. Digital Image Processing (4th Edition)[M]. Pearson, 2018.',
    '[4] Szeliski R. Computer Vision: Algorithms and Applications (2nd Edition)[M]. Springer, 2022.',
    '[5] Microsoft. Asynchronous Programming with async and await[EB/OL]. https://learn.microsoft.com/en-us/dotnet/csharp/asynchronous-programming/',
    '[6] Albahari J, Albahari B. C# 12 in a Nutshell[M]. O\'Reilly Media, 2023.',
    '[7] Microsoft Fluent Design System[EB/OL]. https://fluent2.microsoft.design/',
    '[8] Refactoring UI[EB/OL]. https://www.refactoringui.com/',
    '[9] 字节跳动校园招聘 - Windows客户端开发工程师[EB/OL]. https://jobs.bytedance.com/campus/position/7530950785564231944/detail',
    '[10] Anthropic. Claude Code Documentation[EB/OL]. https://docs.anthropic.com/en/docs/claude-code',
    '[11] Petzold C. Programming Windows (6th Edition)[M]. Microsoft Press, 2013.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

# ============================
# Appendices
# ============================
doc.add_page_break()
add_heading_text('附件一、项目方案文档（Project_Proposal.md）', level=1)
add_para('详见随项目源码一同提交的 Project_Proposal.md 文件。内容包括：项目背景、功能需求分析、技术选型理由、系统架构设计、开发环境说明。', indent=False)
add_para('项目GitHub仓库地址：https://github.com/Xentire/Windows-.git', bold=True, indent=False)

add_heading_text('附件二、测试说明文档（Testing_Report.md）', level=1)
add_para('详见随项目源码一同提交的 Testing_Report.md 文件。内容包括：TDD开发过程、36项测试用例清单（全部通过）、关键测试用例分析及TDD实践效果总结。', indent=False)

# ============================
output_path = r'C:\Users\Xentire\Desktop\Windows程序设计-期末作业\Windows程序设计课程作业-ImageProcessingTool-v2.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
